#!/usr/bin/env python
# encoding=utf-8

#导入依赖的所有库
import os
import shutil
import re
import time
import datetime
import selenium
import docx
import xlrd
import xlwt
import numpy as np
import pandas as pd
import openpyxl as ox
import requests as res

#使用正则表达式筛选pandas.DataFrame的某一列，并返回筛选结果(类似Excel筛选单利)：
def filterdf(regexstr,key,indf,match=False):
    '''
    filter a pandas.DataFrame in one specific column with regular expression.
    key must be in the columns of the DataFrame.
    '''
    import re
    import pandas as pd
    regexstr=re.compile(regexstr)
    if key in indf.columns:
        pass
    else:
        print('key must be in the columns of the DataFrame!')
        return None
    resuli=[]
    for i in indf[key]:
        if match==True:
            reg1=re.match(regexstr,i)
        else:
            reg1=re.search(regexstr,i)
        if reg1 != None:
            resuli.append(i)
    #得到了resuli
    #print(resuli)
    if len(resuli) != 0:
        resurows=[]
        resu=pd.DataFrame(columns=indf.columns)
        for j in resuli:
            a=indf[indf[key]==j]
            #print(a)
            resurows.append(a)
        resu=pd.concat(resurows)
    else:
        resu=pd.DataFrame(columns=indf.columns)
    return resu
#
# Catalogue Test:检测目录是否被别人做了
def catatest(catali,tsl=0,hide=True,pd=False):
        '''
        catalogue test
        pass a 'catalogue list' parameter,return a 'kezuo and yiyou'2-dimension list
        catali must be a list, consisting of the topics that you want to search on the site.
        '''
        import time
        import datetime
        import numpy as np
        import selenium
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        option = webdriver.ChromeOptions()
        option.add_argument("headless")

        t1=datetime.datetime.now()

        if hide==True:
            sch = webdriver.Chrome(options=option)
        else:
            sch = webdriver.Chrome(options=option)
        #sch=webdriver.Chrome()

        site=r'https://workbench.qyresearch.com/login'
        sch.get(site)

        time.sleep(tsl*np.random.rand()+1.002)

        sch.find_element_by_id('username').clear()
        sch.find_element_by_id('username').send_keys('目录查重')
        sch.find_element_by_id('password').clear()
        sch.find_element_by_id('password').send_keys('QYR168168')
        sch.find_element_by_id('password').send_keys(Keys.ENTER)

        time.sleep(tsl*np.random.rand()+1.002)

        kezuo=[]
        yiyou=[]

        n=0
        for i in catali:
            inp=sch.find_element_by_xpath(r'//*[@id="keywords"]')
            inp.send_keys(i)

            time.sleep(tsl*np.random.rand()+1.002)

            sebu=sch.find_element_by_xpath(r'/html/body/app-root/app-layout/div/section/div/app-global-search/div/div[2]/form/div[3]/button')
            sebu.click()

            time.sleep(tsl*np.random.rand()+1.002)

            try:
                a=sch.find_element_by_xpath(r'/html/body/app-root/app-layout/div/section/div/app-global-search/div/app-common-table/div/div[2]')
                if a.text=='暂无数据 No Data':
                    #print(a.text)
                    kezuo.append(i)
                    n+=1
                else:
                    pass
            except:
                b1=sch.find_element_by_xpath(r'/html/body/app-root/app-layout/div/section/div/app-global-search/div/app-common-table/*')
                b2=b1.text
                yiyou.append(b2)
            inp.clear()
        sch.quit()
        #print('可做的哦：','一共',n,'个','\n','----','\n',kezuo)

        if pd==True:
            import pandas as pd
            shuru=pd.Series(catali,name='输入')
            kezuo=pd.Series(kezuo,name='可做')
            resu=pd.concat([shuru,kezuo],axis=1,join='outer')
        else:
            resu=[kezuo,yiyou]

        t2=datetime.datetime.now()
        print('总共用时： ',t2-t1)

        return resu

#接到Sample预处理，在server上搜索节选库和产品库，找到相关文件全部复制：
def samprepare(item,mtch=False):
    '''
    item is a str indicating the name of the study objective.
    '''
    import os
    import re
    import shutil
    import datetime

    t1=datetime.datetime.now()
    #print('start time: ',t1)

    badir=r'D:\HZ.SK\MissionAccomplished\Sample节选'

    #创建item专属的目录
    itemdir=str(badir)+os.sep+str(item)
    os.mkdir(itemdir)

    #创建 Final目录
    final=itemdir+os.sep+str(r'Final')
    os.mkdir(final)

    #在item专属目录下，创建多个参考目录

    reprefdir=itemdir+os.sep+str(r'ReportRef')
    os.mkdir(reprefdir)

    samrefdir=itemdir+os.sep+str(r'SampleRef')
    os.mkdir(samrefdir)

    #contrefdir=itemdir+os.sep+str(r'ContentRef')
    #os.mkdir(contrefdir)

    prodrefdir=itemdir+os.sep+str(r'ProductRef')
    os.mkdir(prodrefdir)



    fdir=[r'\\server\1Report',
         r'\\server\2Sample节选库',
         r'\\server\3Content目录库',
         r'\\server\4Product产品库']

    def findcpall(item,fdir,destination,match=False):
        import os
        import re
        import shutil
        #rs=[]
        for i,j,k in os.walk(fdir):
            for na in k:
                if match==True:
                    regitem=re.compile(item)
                    d=re.match(regitem,na)
                    #if d is not None:
                     #   rsdir=os.path.join(i,na)
                      #  shutil.copy(rsdir,destination)
                else:
                    regitem=re.compile(item)
                    d=re.search(regitem,na)
                if d is not None:
                    rsdir=os.path.join(i,na)
                    shutil.copy(rsdir,destination)

    findcpall(item,fdir[0],reprefdir,match=mtch)
    findcpall(item,fdir[1],samrefdir,match=mtch)
    #findcpall(item,fdir[2],contrefdir,match=mtch)
    findcpall(item,fdir[3],prodrefdir,match=mtch)

    t2=datetime.datetime.now()
    #print('finish time: ',t2)
    print('总共用时： ',t2-t1)
#
def samprepareplus(item,sparams=[r'GIR',None,None],badir=r'D:\HZ.SK\MissionAccomplished\Sample节选',blankdir=r'D:\HZ.SK\模版和操作说明\最新常用Word模板9.12',mtch=False):
    '''
    item is a str indicating the name of the study objective.
    sparams is a list whose length is 3.Pass as much as three parameters and then you get a list of Word-version of blank Samples.
    '''
    import os
    import re
    import shutil
    import datetime

    t1=datetime.datetime.now()
    #print('start time: ',t1)

    #创建item专属的目录
    itemdir=str(badir)+os.sep+str(item)
    os.mkdir(itemdir)
    #创建项目概况说明文件
    txtdir=itemdir+os.sep+str(item)+str(r'.txt')
    #print(txtdir)
    with open(txtdir,'a',encoding='utf-8') as f:
        f.write(item)
        f.write('\n')

    #创建 Final目录
    final=itemdir+os.sep+str(r'Final')
    os.mkdir(final)
    #创建 SampeBlank目录
    sampleblank=itemdir+os.sep+str(r'SampeBlank')
    os.mkdir(sampleblank)

    #在item专属目录下，创建多个参考目录

    #reprefdir=itemdir+os.sep+str(r'ReportRef')
    #os.mkdir(reprefdir)

    samrefdir=itemdir+os.sep+str(r'SampleRef')
    os.mkdir(samrefdir)

    #contrefdir=itemdir+os.sep+str(r'ContentRef')
    #os.mkdir(contrefdir)

    prodrefdir=itemdir+os.sep+str(r'ProductRef')
    os.mkdir(prodrefdir)

    fdir=[r'\\server\1Report',
         r'\\server\2Sample节选库',
         r'\\server\3Content目录库',
         r'\\server\4Product产品库']

    def findcpall(item,fdir,destination,match=False):
        import os
        import re
        import shutil
        #rs=[]
        if match==True:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.match(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        else:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.search(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        #print('copy finished')
    def blankcpall(sparams):
        '''
        sparams is a list whose length is 3.
        Pass as much as three parameters and then you get a list of Word-version of blank Samples.
        '''
        blankli=os.listdir(blankdir)
        def blankdraw(teststr,inli):
            draw=[]
            for i in inli:
                if teststr in i:
                    draw.append(i)
            return draw
        resu1=blankdraw(sparams[0],blankli)
        resu2=blankdraw(sparams[1],resu1)
        resu3=blankdraw(sparams[2],resu2)

        cpli=[]
        for i in resu3:
            a=os.path.join(os.path.abspath(blankdir),i)
            i_fake=re.sub(r'@@@@',item,i,count=1)
            i_fake=re.sub(r'^.+(?=Global)',r'Sample-',i_fake,count=1)
            b=os.path.join(os.path.abspath(sampleblank),i_fake)
            shutil.copy(a,b)
    blankcpall(sparams)

    #findcpall(item,fdir[0],reprefdir,match=mtch)
    findcpall(item,fdir[1],samrefdir,match=mtch)
    #findcpall(item,fdir[2],contrefdir,match=mtch)
    findcpall(item,fdir[3],prodrefdir,match=mtch)

    t2=datetime.datetime.now()
    #print('finish time: ',t2)
    print('总共用时： ',t2-t1)
#
#接到Project预处理，在server上搜索已有报告和sample，找到相关文件全部复制：
def projpp(item,datainfo=True,mtch=False):
    '''
    item is a str indicating the name of the study objective,and do not use regular expression.
    datainfo=True, means you're searching for data information, else you're getting chapter information.
    '''
    import os
    import re
    import shutil
    import datetime

    t1=datetime.datetime.now()
    #print('start time: ',t1)

    if datainfo==True:
        badir=r'D:\HZ.SK\MissionAccomplished\Project 数据'
    else:
        badir=r'D:\HZ.SK\MissionAccomplished\Project 章节'

    #创建item专属的目录
    itemdir=str(badir)+os.sep+str(item)
    os.mkdir(itemdir)

    #创建 Final目录
    final=itemdir+os.sep+str(r'Final')
    os.mkdir(final)

    #在item专属目录下，创建多个参考目录

    reprefdir=itemdir+os.sep+str(r'ReportRef')
    os.mkdir(reprefdir)

    samrefdir=itemdir+os.sep+str(r'SampleRef')
    os.mkdir(samrefdir)

    #contrefdir=itemdir+os.sep+str(r'ContentRef')
    #os.mkdir(contrefdir)

    #prodrefdir=itemdir+os.sep+str(r'ProductRef')
    #os.mkdir(prodrefdir)

    fdir=[r'\\server\1Report',
         r'\\server\2Sample节选库',
         r'\\server\3Content目录库',
         r'\\server\4Product产品库']

    def findcpall(item,fdir,destination,match=False):
        import os
        import re
        import shutil
        #rs=[]
        if match==True:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.match(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        else:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.search(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        #print('copy finished')

    findcpall(item,fdir[0],reprefdir,match=mtch)
    findcpall(item,fdir[1],samrefdir,match=mtch)
    #findcpall(item,fdir[2],contrefdir,match=mtch)
    #findcpall(item,fdir[3],prodrefdir,match=mtch)

    t2=datetime.datetime.now()
    #print('finish time: ',t2)
    print('总共用时： ',t2-t1)
#
#项目预处理-高级版
def projprepareplus(item,sparams=[r'GIR',None,None],datainfo=True,badir=r'D:\HZ.SK\MissionAccomplished',blankdir=r'D:\HZ.SK\模版和操作说明\最新常用Word模板9.12',mtch=False):
    '''
    Pass as much as three parameters confirm the version of the report and then you get a list of Word-version of blank Samples.
    item is a str indicating the name of the study objective.
    sparams is a list whose length is 3.
    datainfo determines which part of work you are taking charge in.
    mtch is the matching type of regular expression (search or match).
    badir,blankdir can be customized. badir is the directory where you put your project files and blankdir is where the blank templates locate.
    '''
    import os
    import re
    import shutil
    import datetime

    if sparams[1]==None:
        sparams[1]=sparams[0]
    if sparams[2]==None:
        sparams[2]=sparams[0]

    if datainfo==True:
        badir=os.path.join(badir,r'Project 数据')
    else:
        badir=os.path.join(badir,r'Project 章节')

    t1=datetime.datetime.now()
    #print('start time: ',t1)

    #创建item专属的目录
    itemdir=str(badir)+os.sep+str(item)
    os.mkdir(itemdir)
    #创建txt项目概况说明文档
    txtdir=itemdir+os.sep+str(item)+str(r'.txt')
    #print(txtdir)
    with open(txtdir,'a',encoding='utf-8') as f:
        f.write(item)
        f.write('\n')

    #创建 Final目录
    final=itemdir+os.sep+str(r'Final')
    os.mkdir(final)
    #创建 SampeBlank目录
    sampleblank=itemdir+os.sep+str(r'SampeBlank')
    os.mkdir(sampleblank)

    #在item专属目录下，创建多个参考目录

    externalRefdir=itemdir+os.sep+str(r'ExternalRef')
    os.mkdir(externalRefdir)

    reprefdir=itemdir+os.sep+str(r'ReportRef')
    os.mkdir(reprefdir)

    samrefdir=itemdir+os.sep+str(r'SampleRef')
    os.mkdir(samrefdir)

    #contrefdir=itemdir+os.sep+str(r'ContentRef')
    #os.mkdir(contrefdir)

    #prodrefdir=itemdir+os.sep+str(r'ProductRef')
    #os.mkdir(prodrefdir)

    fdir=[r'\\server\1Report',
         r'\\server\2Sample节选库',
         r'\\server\3Content目录库',
         r'\\server\4Product产品库']

    def findcpall(item,fdir,destination,match=False):
        import os
        import re
        import shutil
        #rs=[]
        if match==True:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.match(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        else:
            for i,j,k in os.walk(fdir):
                for na in k:
                    regitem=re.compile(item)
                    d=re.search(regitem,na)
                    if d is not None:
                        rsdir=os.path.join(i,na)
                        shutil.copy(rsdir,destination)
        #print('copy finished')
    def blankcpall(sparams):
        if sparams[1]==None:
            sparams[1]=sparams[0]
        if sparams[2]==None:
            sparams[2]=sparams[0]
        blankli=os.listdir(blankdir)
        def blankdraw(teststr,inli):
            draw=[]
            for i in inli:
                if teststr in i:
                    draw.append(i)
            return draw
        resu1=blankdraw(sparams[0],blankli)
        resu2=blankdraw(sparams[1],resu1)
        resu3=blankdraw(sparams[2],resu2)

        cpli=[]
        for i in resu3:
            a=os.path.join(os.path.abspath(blankdir),i)
            i_fake=re.sub(r'@@@@',item,i,count=1)
            i_fake=re.sub(r'^.*(?=[A-Z]{3}.+)',r'章节1-',i_fake,count=1)
            b=os.path.join(os.path.abspath(sampleblank),i_fake)
            shutil.copy(a,b)
    blankcpall(sparams)

    findcpall(item,fdir[0],reprefdir,match=mtch)
    findcpall(item,fdir[1],samrefdir,match=mtch)
    #findcpall(item,fdir[2],contrefdir,match=mtch)
    #findcpall(item,fdir[3],prodrefdir,match=mtch)

    t2=datetime.datetime.now()
    #print('finish time: ',t2)
    print('总共用时： ',t2-t1)
#

# Find files with Regular Expression：正则表达式搜索本地文件
def findfile(item,fdir,match=False):
    '''
    Find the file you need in a directory.
    Regular Expression is supported.
    '''
    import os
    import re
    rs=[]
    for i,j,k in os.walk(fdir):
        for na in k:
            if match==True:
                regitem=re.compile(item)
                d=re.match(regitem,na)
                if d is not None:
                    rs.append(os.path.join(i,na))
            else:
                regitem=re.compile(item)
                d=re.search(regitem,na)
                if d is not None:
                    rs.append(os.path.join(i,na))
    if rs==[]:
        print(r"Can't find anything!")
    return rs

#批量复制文件,其实没用,先留着吧.
def copyall(cpli,destination):
    import shutil
    for i in cpli:
        shutil.copy(i,destination)
#
#搜索文件，并复制所有搜索结果到指定文件夹
def findcpall(item,fdir,destination,match=False):
    import os
    import re
    import shutil
    #rs=[]
    for i,j,k in os.walk(fdir):
        for na in k:
            if match==True:
                regitem=re.compile(item)
                d=re.match(regitem,na)
                #if d is not None:
                 #   rsdir=os.path.join(i,na)
                  #  shutil.copy(rsdir,destination)
            else:
                regitem=re.compile(item)
                d=re.search(regitem,na)
            if d is not None:
                rsdir=os.path.join(i,na)
                shutil.copy(rsdir,destination)
    #return rs

#用正则表达式匹配列表中的元素,返回符合条件的元素列表：
def regtestli(item,li,match=False):
    '''
    Test elements in a list with Regular Expression.
    Extract elements that work well with a Regular Expression (item) from a list(li), and return a result list.
    '''
    import re
    c=re.compile(item)
    n=0
    rs=[]
    for i in li:
        if match == True:
            b=re.match(c,i)
            if b is not None:
                rs.append(li[n])
        else:
            b=re.search(c,i)
            if b is not None:
                rs.append(li[n])
        n+=1
    return rs

#打开若干网站,使用selenium.
def opsite(siteli,tsl=0):
    '''
    open a list of websites.
    siteli must be a list, each element of which is one of the websites that you want to open in one single web browser.
    '''
    import time
    import numpy as np
    from selenium import webdriver
    osite=webdriver.Chrome()
    n=1
    for i in siteli:
        osite.get(i)
        time.sleep(tsl*np.random.rand()+0.002)
        osite.execute_script("window.open()")
        handles=osite.window_handles
        osite.switch_to.window(handles[n])
        n+=1
    osite.close()
    osite.switch_to.window(handles[0])

#简单Google搜索：按照对象列表com+关键词列表kw组合搜索。使用selenium打开网站.
def simplegoo(com,kw,tsl=0):
    '''
    Simply Google (simplegoo) companies plus the keyword given.
    com and kw must be lists, excluding NaN elements.
    tsl is the time sleep level when opening each tab.
    '''
    import time
    import numpy as np
    from selenium import webdriver
    from selenium.webdriver.common.keys import Keys

    wh=r"https://www.google.com/"
    js=r"window.open('https://www.google.com/')"

    sch=webdriver.Chrome()

    li=[]
    for i in com:
        for j in kw:
            b=str(i)+r' '+str(j)
            li.append(b)
    n=1
    for i in li:
            a=i
            sch.get(wh)
            sch.find_element_by_name("q").send_keys(a)
            sch.find_element_by_name("q").send_keys(Keys.ENTER)
            time.sleep(tsl*np.random.rand()+0.002)
            sch.execute_script(js)
            handles=sch.window_handles
            sch.switch_to.window(handles[n])
            n+=1
    sch.close()
    sch.switch_to.window(handles[0])
#

#自动填入ProductName和CompanyList
#做目录的Excel里边不同的产品行数相差56。
#产品中文名称：B2；Product Name：C2（c58） #company 1 的位置：C4(c60)
def fillcali(cali,indir,savedir,start_prod='C2',start_com='C4',cellstep=56):
    '''
    Fill the Category List (fillcali) into the Excel.
    catali must be a pandas.DataFrame where the first product starts at the second column,
    and of course, the first column would be the "id" column.
    '''
    import openpyxl as ox
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter, column_index_from_string

    wb=ox.load_workbook(indir)
    ws=wb['产品信息']

    nali=cali.columns
    nali=list(nali)
    nali.pop(0)

    def onefill(prodbcell,combcell,n):
        '''
        n的范围从0开始，表示列表nali当中元素的索引。
        n starts from 0, indicating to the index of the list nali.
        '''
        prodbcell.value=nali[n]

        comli=cali[nali[n]]
        comli=comli.iloc[range(12)].dropna()
        comli=list(comli)
        le=len(comli)
        for i in range(le):
            combcell.offset(i,0).value=comli[i]

    prodbcell=ws[start_prod]
    combcell=ws[start_com]

    for i in range(len(nali)):
        onefill(prodbcell,combcell,i)
        prodbcell=prodbcell.offset(cellstep,0)
        combcell=combcell.offset(cellstep,0)

    wb.save(savedir)
#

#函数-获取docx文档概览：
def getscan(indir,depth=1,accu=True):
    '''
    Get scan of a Microsoft Word file.
    indir is the inputing directory,and depth indicates how deep you want to know.
    e.g.with depth=3, you'll get all the Heading 1, Heading 2 as well as Heading 3 of the file.
    and,if accu is False, you'll get Heading 3 only.
    '''
    resu=[]
    import docx
    f=docx.Document(indir)
    for i in f.paragraphs:
        if accu==False:
            name_left='Heading '
            name_right=str(depth)
            if i.style.name==name_left+name_right:
                resu.append(i.text)
        else:
            j=1
            for j in range(1,depth+1):
                name_left='Heading '
                name_right=str(j)
                if i.style.name==name_left+name_right:
                    resu.append(i.text)
                    j+=1
    return resu
#
#把文档中所有的paragraph作为list中的element组成list：
def transferli(indir):
    import docx
    f=docx.Document(indir)
    resu=[]
    for i in f.paragraphs:
        resu.append(i.text)
    return resu
#
#获取Word文档指定章节的内容：
#def

#对比两个Word文档(docx)结构的不同，返回一个pandas.DataFrame:
def dcompare(sdir1,sdir2,dep=1,accum=True):
    def getscan(indir,depth=1,accu=True):
        resu=[]
        import docx
        f=docx.Document(indir)
        for i in f.paragraphs:
            if accu==False:
                name_left='Heading '
                name_right=str(depth)
                if i.style.name==name_left+name_right:
                    resu.append(i.text)
            else:
                j=1
                for j in range(1,depth+1):
                    name_left='Heading '
                    name_right=str(j)
                    if i.style.name==name_left+name_right:
                        resu.append(i.text)
                        j+=1
        return resu

    s1=getscan(sdir1,depth=dep,accu=accum)
    s2=getscan(sdir2,depth=dep,accu=accum)
    import pandas as pd
    s1=pd.Series(s1)
    s2=pd.Series(s2)
    dresu=pd.concat([s1,s2],axis=1,join='outer')
    return dresu
#

#get structure of a word template:
def getStructure(proj_name,reading_dir,outputing_dir):
    '''
    get the data structure of a word template/sample/report.
    '''
    subj=pd.Series(['[Tt]op\\s*\\d*','[Mm]arket\\s*[Ss]ize|[Mm]arket\\s*[Vv]alue','[Rr]evenue','[Ss]ales?','[Cc]apacity','[Pp]roduction','[Cc]onsumption','[Ee]xport','[Ii]mport','[Ss]ize','[Vv]alue','[Pp]rice'])
    shareif=pd.Series([r'[Ss]hare'])
    bywhat=pd.Series(['[Bb]y\\s*[Tt]ypes?','[Bb]y\\s*[Pp]roduct','[Bb]y\\s*[Aa]pplications?','[Bb]y\\s*[Rr]egions?','[Bb]y\\s*[Cc]ountr[y|(ies)]','[Bb]y\\s*(([Cc]ompan[y|(ies)])|([Pp]layers?)|([Mm]anufacturers?))'])
    rangeyear=pd.Series([r'\d{4}(\s*(to|-|-)\s*\d{4})?'])
    regionli=pd.Series(['[Gg]lobal','[Nn]orth\\s*America','[Ee]urope','[Aa]sia(\\s*|-)Pacific','[Ss]outh\\s*America','[Mm]iddle\\s*[Ee]ast(\\s*and\\s*[Aa]frica)?','Americas?','United\\s*States(\\s*of\\s*America)?'])
    countryli=pd.Series(['USA','Canada','Mexico','Brazil','China','Taiwan','Japan','(South\\s*)?Korea','Southeast\\s*Asia','India','Australia','Germany','France','UK','Italy','Russia','Spain','Egypt','South\\s*Africa','Israel','Turkey','GCC\\s*Countries'])
    vardict=pd.concat([subj,shareif,bywhat,rangeyear,regionli,countryli],axis=1,join='outer')
    vardict.columns=['subj','shareif','bywhat','rangeyear','regionli','countryli']

    out_name=r'SimpleInfo-'+proj_name+r'.xlsx'
    outdir_fake=os.path.join(outdir_fake1,out_name)
    print(r'outputing directory:',outputing_dir)

    def readstrli(instrli,vardict=pd.read_excel(r'D:\HZ.SK\MissionAccomplished\Project 数据\vardict-whatdata.xlsx',sheet_name=r'vardict')):
        ##
        def readstr(instr):
            def oneread(tli,tstr):
                bli=[]
                for i in tli:
                    acom=re.compile(i)
                    a=re.search(acom,tstr)
                    if a != None:
                        b=a.group(0)
                        bli.append(b)
                    else:
                        bli.append('nothing')
                for j in bli:
                    if j != 'nothing':
                        b_final=j
                        return b_final
            vardict_col=list(vardict.columns)
            subinfo = []
            for i in vardict_col:
                tli1=vardict.loc[:,i].dropna()
                tli1=list(tli1)
                c=oneread(tli1,instr)
                subinfo.append(c)
            return subinfo
            print(subinfo)
        #
        infomat=[]
        for i in instrli:
            i_fake=str(i)
            subinfo1=readstr(i_fake)
            infomat.append(subinfo1)
        infomat=pd.DataFrame(infomat,columns=vardict.columns)
        return infomat
    #
    def drawtf(indir):
        from docx import Document
        indoc=Document(indir)
        tfli=[]
        for i in indoc.paragraphs:
            if i.style.name=='Table Style':
                tfli.append(r'Table '+i.text)
            elif i.style.name=='Figure Style':
                tfli.append(r'Figure '+i.text)
            else:
                pass
        return tfli
    #
    t1=drawtf(reading_dir)
    print(r'图表数目：',len(t1))
    t3=readstrli(t1)
    t3['original']=t1
    t3.to_excel(outdir_fake,sheet_name=r'Catalogue')
#


#定义Sch类
#简单Google搜索
from selenium import webdriver
class Sch(webdriver.Chrome):
    #简单搜索
    def simplegoo(self,com,kw=[''],tsl=0):
        '''
        Google companies plus the keyword given.
        com and kw must be lists, excluding NaN elements.
        tsl is the time sleep level when opening each tab.
        '''
        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r"https://www.google.com/"
        js=r"window.open('https://www.google.com/')"

        li=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                li.append(b)
        n=1
        for i in li:
                a=i
                self.get(wh)
                self.find_element_by_name("q").send_keys(a)
                self.find_element_by_name("q").send_keys(Keys.ENTER)
                time.sleep(tsl*np.random.rand()+0.002)
                self.execute_script(js)
                handles=self.window_handles
                self.switch_to.window(handles[n])
                n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    #快速搜搜
    def fastgoo(self,com,kw=[''],tsl=0):
        qkw=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                qkw.append(b)
        #
        siteli=[]
        for i in qkw:
            r1=res.get(r'https://www.google.com/search?source',params={'q':i,'client':'firefox-b-d'})
            r2=r1.url
            siteli.append(r2)
        n=1
        for i in siteli:
            self.get(i)
            self.execute_script("window.open()")
            handles=self.window_handles
            self.switch_to.window(handles[n])
            n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    def simpledu(self,com,kw=[''],tsl=0):
        '''
        Google废了的时候用百度，比必应还厉害呢！
        '''
        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r"https://www.baidu.com/"
        js=r"window.open('https://www.baidu.com/')"

        li=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                li.append(b)
        n=1
        for i in li:
                a=i
                self.get(wh)
                self.find_element_by_id("kw").send_keys(a)
                self.find_element_by_id("kw").send_keys(Keys.ENTER)

                time.sleep(tsl*np.random.rand()+0.502)

                self.execute_script(js)
                handles=self.window_handles
                self.switch_to.window(handles[n])
                n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    #简单百度
    def simplebi(self,com,kw=[''],tsl=0):
        '''
        其实必应搜索并没有百度厉害
        '''
        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r"https://bing.com/"
        js=r"window.open('https://bing.com/')"

        li=[]
        for i in com:
            for j in kw:
                b=str(i)+r' '+str(j)
                li.append(b)
        n=1
        for i in li:
                a=i
                self.get(wh)
                self.find_element_by_id("sb_form_q").send_keys(a)
                self.find_element_by_id("sb_form_q").send_keys(Keys.ENTER)
                time.sleep(tsl*np.random.rand()+0.45)
                self.find_element_by_id("est_en").click()
                time.sleep(tsl*np.random.rand()+0.2)
                self.execute_script(js)
                handles=self.window_handles
                self.switch_to.window(handles[n])
                n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    #简单搜Wikipedia
    def wiki(self,li,tsl=0):

        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r'https://en.wikipedia.org/wiki/Main_Page'
        js=r"window.open('https://en.wikipedia.org/wiki/Main_Page')"

        n=1
        for i in li:
            a=str(i)
            self.get(wh)
            inp=self.find_element_by_id("searchInput")
            inp.send_keys(a)
            cl=self.find_element_by_id(r'searchButton')
            cl.click()
            time.sleep(tsl*np.random.rand()+0.45)
            self.execute_script(js)
            handles=self.window_handles
            self.switch_to.window(handles[n])
            n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    #search on Gartner
    def schgtn(self,catali,tsl=0):

        import time
        import numpy as np
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r'https://www.gartner.com/reviews/home'
        js=r'window.open("https://www.gartner.com/reviews/home")'

        links=[]

        n=1
        for i in catali:
            a=str(i)
            self.get(wh)
            time.sleep(tsl*np.random.rand()+0.2)
            #//*[@id="app"]/div/div[1]/div[4]/div/div/div[1]/div/div/div/div[1]/label/input
            inp=self.find_element_by_xpath('//*[@id="app"]/div/div[1]/div[4]/div/div/div[1]/div/div/div/div[1]/label/input')
            inp.send_keys(a)
            time.sleep(tsl*np.random.rand()+0.8)

            inp.send_keys(Keys.ENTER)
            time.sleep(tsl*np.random.rand()+0.12)

            cururl=self.current_url
            links.append(cururl)

            self.execute_script(js)
            handles=self.window_handles
            self.switch_to.window(handles[n])
            n+=1
        self.close()
        self.switch_to.window(handles[0])
        return links
    #
    def getvdgtn(self,catali,tsl=0,wtint=7):

        import time
        import numpy as np
        import pandas as pd
        from selenium import webdriver
        from selenium.webdriver.common.keys import Keys

        wh=r'https://www.gartner.com/reviews/home'
        js=r'window.open("https://www.gartner.com/reviews/home")'

        def getvenli(self):
            venli=[]
            fvendors=self.find_elements_by_css_selector(r'.col-md-6 div a')
            for i in fvendors:
                venli.append(i.text)
            return venli

        resu=pd.DataFrame([])

        for i in catali:
            a=str(i)
            self.get(wh)
            time.sleep(tsl*np.random.rand()+0.2)
            inp=self.find_element_by_xpath('//*[@id="app"]/div/div[1]/div[4]/div/div/div[1]/div/div/div/div[1]/label/input')
            inp.send_keys(a)
            time.sleep(tsl*np.random.rand()+0.8)

            inp.send_keys(Keys.ENTER)
            time.sleep(tsl*np.random.rand()+1.2)
            time.sleep(tsl*np.random.rand()+wtint)
            #try:
            showvend=self.find_element_by_xpath(r'//*[@id="app"]/div/div[1]/div[4]/div/div[2]/div/section/div/div[1]/div[1]/div/div[2]/a')
            showvend.click()
            #except:
             #   pass
            vdlist=getvenli(self)
            vdlist=pd.Series(vdlist,name=a)
            resu=pd.concat([resu,vdlist],axis=1,join='outer')
        self.close()
        return resu
    #
    #打开若干网站
    def opsite(self,siteli,tsl=0):
        '''
        open a list of websites.
        siteli must be a list, each element of which is one of the websites that you want to open in one single web browser.
        '''
        import time
        import numpy as np
        from selenium import webdriver

        n=1
        for i in siteli:
            self.get(i)
            time.sleep(tsl*np.random.rand()+0.002)
            self.execute_script("window.open()")
            handles=self.window_handles
            self.switch_to.window(handles[n])
            n+=1
        self.close()
        self.switch_to.window(handles[0])
    #
    def goto(self,one_site):
        js=r"window.open("+r'"'+one_site+r'")'
        #print(js)
        self.execute_script(js)
    #
    def gotomany(self,siteli):
        for i in siteli:
            self.goto(i)
        handles=self.window_handles
        self.switch_to.window(handles[0])
        self.close()
        self.switch_to.window(handles[0])
    #
    #Catalogue Test
    #def
#
